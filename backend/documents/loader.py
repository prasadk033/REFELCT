"""
Document loading and text extraction.

Supports: PDF, DOCX, TXT
Uses Haystack for PDF/TXT, python-docx for DOCX.
Extracts embedded images from PDFs for TurboOCR processing.
"""
import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Tuple

from haystack import Document
from haystack.components.converters import PyPDFToDocument, TextFileToDocument
import docx

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Loads and extracts text from PDF, DOCX, and TXT documents."""

    def __init__(self):
        self.pdf_converter = PyPDFToDocument()
        self.txt_converter = TextFileToDocument()

    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document and extract text content.

        Args:
            file_path: Absolute path to the document.

        Returns:
            List of Haystack Document objects with extracted text.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        ext = path.suffix.lower()

        if ext == '.pdf':
            try:
                result = self.pdf_converter.run(sources=[str(path)])
                docs = result.get("documents", [])
                if docs and any(d.content and d.content.strip() for d in docs):
                    return docs
            except Exception as e:
                logger.warning(f"PyPDF conversion failed ({e}). Attempting text fallback.")

            # Fallback for plain-text or non-standard documents
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                if content.strip():
                    return [Document(content=content, meta={"file_path": str(path), "type": "pdf_text_fallback"})]
            except Exception:
                pass
            return []

        elif ext == '.txt':
            result = self.txt_converter.run(sources=[str(path)])
            return result["documents"]

        elif ext in ('.docx', '.doc'):
            return self._load_docx(str(path))

        elif ext in ('.jpg', '.jpeg', '.png', '.webp'):
            return self._load_image(str(path))

        else:
            raise ValueError(f"Unsupported document type: {ext}. Supported types: PDF, TXT, DOCX, JPG, PNG.")

    def _load_image(self, file_path: str) -> List[Document]:
        """Extract text from an image using TurboOCR."""
        from documents.turboocr import turbo_ocr
        path = Path(file_path)
        extracted = ""
        try:
            with open(path, "rb") as f:
                img_bytes = f.read()
            if turbo_ocr.is_available:
                res = turbo_ocr.extract_from_image(img_bytes, path.name)
                if res.get("success") and res.get("text"):
                    extracted = res["text"]
                elif res.get("error"):
                    logger.warning(f"TurboOCR extraction error on {path.name}: {res['error']}")
        except Exception as e:
            logger.warning(f"Image OCR failed for {file_path}: {e}")

        if not extracted or not extracted.strip():
            extracted = f"[Image Source: {path.name} — Visual architectural reference (No text detected)]"

        return [Document(content=extracted, meta={"file_path": file_path, "type": "image"})]



    def _load_docx(self, file_path: str) -> List[Document]:
        """Extract text from a DOCX file using python-docx."""
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        text_content = "\n".join(full_text)
        return [Document(content=text_content, meta={"file_path": file_path, "type": "docx"})]

    def extract_images_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract embedded images from a PDF file.

        Returns:
            List of dicts: [{"data": bytes, "filename": str, "page": int}]
        """
        images = []
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_images = page.images
                    if page_images:
                        for img_idx, img_info in enumerate(page_images):
                            try:
                                # Extract the image from the page
                                # pdfplumber provides image metadata; extract via the page
                                img = page.crop((
                                    img_info.get("x0", 0),
                                    img_info.get("top", 0),
                                    img_info.get("x1", page.width),
                                    img_info.get("bottom", page.height),
                                ))
                                # Convert cropped region to image
                                pil_image = img.to_image(resolution=200)
                                import io
                                img_bytes = io.BytesIO()
                                pil_image.save(img_bytes, format="PNG")
                                img_data = img_bytes.getvalue()

                                images.append({
                                    "data": img_data,
                                    "filename": f"page{page_num}_img{img_idx}.png",
                                    "page": page_num,
                                })
                            except Exception as e:
                                logger.warning(f"Failed to extract image {img_idx} from page {page_num}: {e}")
                                continue

        except ImportError:
            logger.warning("pdfplumber not available for image extraction")
        except Exception as e:
            logger.error(f"Failed to extract images from PDF: {e}")

        logger.info(f"Extracted {len(images)} images from {file_path}")
        return images

    def extract_text_combined(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Extract text content and identify images for OCR processing.

        Returns:
            Tuple of (extracted_text, list_of_images)
        """
        docs = self.load_document(file_path)
        text = "\n\n".join([d.content for d in docs if d.content])

        # Extract images from PDFs
        images = []
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            images = self.extract_images_from_pdf(file_path)

        return text, images
